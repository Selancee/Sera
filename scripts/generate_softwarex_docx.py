#!/usr/bin/env python3
"""Generate a single-column, line-numbered SoftwareX review DOCX from Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "paper" / "softwarex" / "manuscript" / "seraedit_softwarex.md"
DEFAULT_OUTPUT = ROOT / "paper" / "softwarex" / "manuscript" / "seraedit_softwarex.docx"


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def enable_line_numbers(document: Document) -> None:
    for section in document.sections:
        sect_pr = section._sectPr
        line_numbers = sect_pr.find(qn("w:lnNumType"))
        if line_numbers is None:
            line_numbers = OxmlElement("w:lnNumType")
            sect_pr.append(line_numbers)
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:restart"), "newPage")
        line_numbers.set(qn("w:distance"), "360")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.05
    for name, size in (("Title", 16), ("Heading 1", 13), ("Heading 2", 11.5)):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def add_table(document: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            value = clean_inline(row[column_index]) if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True


def build_document(source: Path, output: Path) -> None:
    document = Document()
    style_document(document)
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("| "):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].startswith("|"):
                cells = [item.strip() for item in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", item) for item in cells):
                    rows.append(cells)
                index += 1
            if rows:
                add_table(document, rows)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(clean_inline(line[2:]), style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            document.add_heading(clean_inline(line[3:]), level=1)
        elif line.startswith("### "):
            document.add_heading(clean_inline(line[4:]), level=2)
        elif line.startswith("!["):
            image_match = re.match(r"!\[[^]]*\]\(([^)]+)\)", line)
            if image_match:
                image_path = (source.parent / image_match.group(1)).resolve()
                if image_path.suffix.lower() == ".svg" and image_path.with_suffix(".png").is_file():
                    image_path = image_path.with_suffix(".png")
                if image_path.is_file():
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.add_run().add_picture(str(image_path), width=Inches(6.1))
        elif line.startswith("- "):
            document.add_paragraph(clean_inline(line[2:]), style="List Bullet")
        else:
            paragraph_lines = [line]
            index += 1
            while index < len(lines):
                next_line = lines[index].rstrip()
                if not next_line or next_line.startswith(("#", "|", "![", "- ")):
                    break
                paragraph_lines.append(next_line)
                index += 1
            paragraph = document.add_paragraph(clean_inline(" ".join(paragraph_lines)))
            if paragraph.text.startswith("Keywords:"):
                paragraph.runs[0].bold = True
            continue
        index += 1

    enable_line_numbers(document)
    properties = document.core_properties
    properties.title = "SeraEdit: Reliable Language-Guided MusicXML Editing through Structured Score Patches"
    properties.subject = "SoftwareX Original Software Publication draft"
    properties.keywords = "MusicXML, symbolic music, score editing, structured patches, validation, research software"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate the line-numbered SoftwareX DOCX review manuscript.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser


def main() -> int:
    args = build_parser().parse_args()
    build_document(args.source, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
