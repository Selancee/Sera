#!/usr/bin/env python3
"""Build editable Word attachments for the SoftwareX submission package."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "paper" / "softwarex" / "submission"
SOURCES = (
    "COVER_LETTER.md",
    "DECLARATION_OF_INTEREST.md",
    "CREDIT_AUTHOR_STATEMENT.md",
    "GENERATIVE_AI_DISCLOSURE.md",
    "DATA_AND_CODE_AVAILABILITY.md",
)


def markdown_blocks(markdown: str) -> list[tuple[str, str]]:
    """Parse the small heading/paragraph/list subset used by submission statements."""
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append(("paragraph", " ".join(paragraph)))
            paragraph.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush()
            blocks.append((f"heading{len(heading.group(1))}", heading.group(2)))
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            flush()
            blocks.append(("bullet", bullet.group(1)))
            continue
        paragraph.append(line.replace("`", ""))
    flush()
    return blocks


def render_markdown_to_docx(source: Path, target: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for kind, text in markdown_blocks(source.read_text(encoding="utf-8")):
        if kind.startswith("heading"):
            level = min(3, int(kind[-1]))
            document.add_heading(text, level=level)
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)

    document.core_properties.title = source.stem.replace("_", " ").title()
    document.core_properties.author = "Yuan Gao"
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def validate_docx(path: Path) -> None:
    document = Document(path)
    visible = " ".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    if len(visible) < 20:
        raise ValueError(f"Generated attachment is unexpectedly empty: {path}")


def build_all(check_only: bool = False) -> list[Path]:
    outputs: list[Path] = []
    for name in SOURCES:
        source = SUBMISSION_DIR / name
        target = source.with_suffix(".docx")
        if not check_only:
            render_markdown_to_docx(source, target)
        validate_docx(target)
        outputs.append(target)
    return outputs


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--check", action="store_true", help="Validate existing DOCX attachments without rebuilding them.")
    args = parser.parse_args()
    for path in build_all(check_only=args.check):
        print(path.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
